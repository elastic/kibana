#!/usr/bin/env python3
"""Generate PLATFORM_GAPS_REQUIREMENTS.docx for SDLC Fleet package platform asks."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = Path(__file__).with_name('PLATFORM_GAPS_REQUIREMENTS.docx')

GAPS = [
    # Fleet
    {
        'id': 'FLEET-001',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P1',
        'title': 'Generic manifest-driven workflow placeholder substitution',
        'requirement': 'Integration packages declare manifest vars; any workflow YAML placeholder `REPLACE_WITH_{VAR}` is substituted at install from the integration policy without Fleet core hardcoding package-specific var names.',
        'current': 'Substitution in `step_install_workflow_assets.ts` hardcodes SDLC var names (github_connector_id, slack_connector_id, etc.).',
        'gap': 'Not reusable for other integrations; every new var requires a Fleet code change.',
        'deliverable': 'Generic substitution map from manifest `vars[]` + convention-based placeholders; deprecate package-specific constants.',
        'acceptance': 'Second integration package substitutes custom vars with zero Fleet code changes; unit tests cover missing/partial vars.',
        'owner': '@elastic/fleet',
        'refs': 'step_install_workflow_assets.ts',
    },
    {
        'id': 'FLEET-002',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P1',
        'title': 'Materialize alerting rules from integration package templates',
        'requirement': 'On Fleet install, integration packages ship `kibana/alerting_rule_template/*.json` and platform creates disabled alerting rules (or registers enablement UI) — not only for `elastic_agent`.',
        'current': '`stepCreateAlertingAssets` returns early for all packages except `elastic_agent`; SDLC templates import as saved objects only.',
        'gap': 'Admins must manually create rules from templates; packaged “Fleet → Integration → Alerting” UX is incomplete.',
        'deliverable': 'Opt-in manifest flag (e.g. `create_alerting_rules: true`) + rule materialization loop for integration packages; rules created disabled.',
        'acceptance': 'SDLC package install creates 4 disabled rules from templates; uninstall removes them; elastic_agent behavior unchanged.',
        'owner': '@elastic/fleet + @elastic/alerting',
        'refs': 'step_create_alerting_assets.ts; PLATFORM_RFC_KIBANA_ASSET_EXTENSIONS.md P4',
    },
    {
        'id': 'FLEET-003',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Index alias assets in Fleet packages',
        'requirement': 'Packages declare Elasticsearch index aliases (e.g. legacy sync-state alias) as first-class install/uninstall assets.',
        'current': 'Index templates and ILM supported; alias creation is manual or custom bootstrap.',
        'gap': 'No `elasticsearch/index_alias/` (or equivalent) asset type in EPM install pipeline.',
        'deliverable': 'New Elasticsearch asset type + install/remove steps; documented in Fleet package spec.',
        'acceptance': 'Package ships alias asset; install creates alias; uninstall removes alias when safe.',
        'owner': '@elastic/fleet',
        'refs': 'PLATFORM_GAPS_RFC F-2',
    },
    {
        'id': 'FLEET-004',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Re-apply workflow/agent substitution on policy var update',
        'requirement': 'When an integration policy var changes (connector ID, org login), package-managed workflows and agents are updated in place.',
        'current': 'Substitution runs at install time only; policy updates do not refresh workflow YAML.',
        'gap': 'Connector rotation requires manual workflow edit or reinstall.',
        'deliverable': 'Hook on package policy update → re-run substitution + createOrUpdate for affected assets.',
        'acceptance': 'Change github_connector_id on policy → all fleet workflows reflect new ID without reinstall.',
        'owner': '@elastic/fleet',
        'refs': 'step_install_workflow_assets.ts; resolvePackagePolicyConnectorVars',
    },
    {
        'id': 'FLEET-005',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Opt-in enable package workflows on install',
        'requirement': 'Manifest option to install workflows as enabled (or enable a defined subset) after connector vars are resolved.',
        'current': 'All SDLC workflows ship `enabled: false`; admin must enable ~25 workflows manually.',
        'gap': 'Poor post-install UX for Kibana-only integrations whose primary deliverable is scheduled ingest.',
        'deliverable': 'Manifest flag `workflows.default_enabled: true|false|list`; respect missing placeholder guard.',
        'acceptance': 'Package with flag + resolved vars installs with workflows enabled; missing placeholders keep disabled + warning.',
        'owner': '@elastic/fleet + @elastic/workflows-eng',
        'refs': 'sdlc_intel manifest.yml; kibana/workflow/*.yaml',
    },
    {
        'id': 'FLEET-006',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Package-managed asset immutability in Kibana UI',
        'requirement': 'Fleet-installed workflows and agents are visibly package-managed; edits warn or are read-only unless detached.',
        'current': 'Assets appear as normal workflows/agents; users can break package upgrades.',
        'gap': 'No generic “managed by integration X” UX contract.',
        'deliverable': 'Metadata flag on workflow/agent SO + UI badge + upgrade reconcile behavior documented.',
        'acceptance': 'Fleet workflow shows managed badge; upgrade overwrites user edits per policy.',
        'owner': '@elastic/fleet + @elastic/workflows-eng + @elastic/agent-builder',
        'refs': 'PLATFORM_RFC open questions §1',
    },
    {
        'id': 'FLEET-007',
        'component': 'Fleet',
        'status': 'Partial',
        'priority': 'P2',
        'title': 'Install workflow/agent assets without HTTP request context',
        'requirement': 'Package install paths that lack KibanaRequest (API/background) can still install workflow and agent assets.',
        'current': 'Both steps skip when `context.request` is missing.',
        'gap': 'Headless or programmatic installs omit primary package value.',
        'deliverable': 'Internal system-user request or request-less code paths on management APIs.',
        'acceptance': 'Install via supported API without user session still creates workflow + agent assets.',
        'owner': '@elastic/fleet',
        'refs': 'step_install_workflow_assets.ts; step_install_agent_assets.ts',
    },
    {
        'id': 'FLEET-008',
        'component': 'Fleet',
        'status': 'Done',
        'priority': '—',
        'title': 'KibanaAssetType.workflow package install/uninstall',
        'requirement': 'Ship workflow YAML under `kibana/workflow/`; install with stable fleet-prefixed IDs; uninstall cleanup.',
        'current': 'Implemented via stepInstallWorkflowAssets + remove.ts.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'SDLC package installs 25+ workflows; uninstall removes them.',
        'owner': '@elastic/fleet',
        'refs': 'step_install_workflow_assets.ts',
    },
    {
        'id': 'FLEET-009',
        'component': 'Fleet',
        'status': 'Done',
        'priority': '—',
        'title': 'KibanaAssetType.agent package install/uninstall',
        'requirement': 'Ship agent YAML under `kibana/agent/`; create persisted Agent Builder agents; uninstall cleanup.',
        'current': 'Implemented via stepInstallAgentAssets + deletePackageManagedAgent.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'SDLC agents install before workflows; agent-id placeholders resolve.',
        'owner': '@elastic/fleet + @elastic/agent-builder',
        'refs': 'step_install_agent_assets.ts; PLATFORM_RFC P1/P2',
    },
    {
        'id': 'FLEET-010',
        'component': 'Fleet',
        'status': 'Done',
        'priority': '—',
        'title': 'Integration knowledge base from package docs',
        'requirement': 'Index `docs/**/*.md` into integration knowledge base on install.',
        'current': 'step_save_knowledge_base indexes markdown from package archive.',
        'gap': 'Requires Enterprise license + Integrations knowledge setting enabled.',
        'deliverable': '—',
        'acceptance': 'SDLC knowledge_base/*.md searchable after install when licensed.',
        'owner': '@elastic/fleet',
        'refs': 'step_save_knowledge_base.ts',
    },
    {
        'id': 'FLEET-011',
        'component': 'Fleet',
        'status': 'Done',
        'priority': '—',
        'title': 'ES|QL view assets in Fleet packages',
        'requirement': 'Package ships `elasticsearch/esql_view/*.yml`; install/uninstall with package.',
        'current': 'step_install_esql_views + cleanup on remove.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'SDLC ES|QL views available to dashboards and agents post-install.',
        'owner': '@elastic/fleet',
        'refs': 'step_install_esql_views.ts',
    },
    {
        'id': 'FLEET-012',
        'component': 'Fleet',
        'status': 'Partial',
        'priority': 'P3',
        'title': 'Package upgrade reconciliation for workflows and agents',
        'requirement': 'Package upgrade updates workflow/agent definitions idempotently without duplicate IDs or stale steps.',
        'current': 'createOrUpdate on install; upgrade edge cases under-tested.',
        'gap': 'No documented upgrade contract for package authors.',
        'deliverable': 'Upgrade semantics doc + tests for version bump over existing fleet assets.',
        'acceptance': 'Upgrade sdlc_intel 0.1.0→0.2.0 updates workflows in place; schedules preserved per policy.',
        'owner': '@elastic/fleet',
        'refs': 'workflowsManagement.createOrUpdateWorkflow',
    },
    {
        'id': 'FLEET-013',
        'component': 'Fleet',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Kibana-only integration first-class UX',
        'requirement': 'Integrations with `inputs: []` (no Elastic Agent) have clear Fleet UI: connector setup checklist, workflow status, no agent policy confusion.',
        'current': 'Works technically; UX reads like traditional agent integration.',
        'gap': 'Platform content and UI do not distinguish Kibana-only ETL integrations.',
        'deliverable': 'Fleet UI pattern for agentless integrations + manifest hint.',
        'acceptance': 'SDLC install wizard shows connector checklist and workflow enablement without agent steps.',
        'owner': '@elastic/fleet',
        'refs': 'sdlc_intel manifest policy_templates',
    },
    # Workflows
    {
        'id': 'WF-001',
        'component': 'Workflows',
        'status': 'Done',
        'priority': '—',
        'title': 'Stock-step ETL without custom step types',
        'requirement': 'Multi-index scheduled ETL expressible with connectors + elasticsearch.* + flow control only.',
        'current': 'SDLC package validates with zero product-specific step types.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'All sdlc_intel workflows use platform step types only.',
        'owner': '@elastic/workflows-eng',
        'refs': 'kibana/workflow/*.yaml',
    },
    {
        'id': 'WF-002',
        'component': 'Workflows',
        'status': 'Done',
        'priority': '—',
        'title': 'elasticsearch.index upsert ingest step',
        'requirement': 'Idempotent document write by `_id` from workflow foreach loops.',
        'current': 'Used across all catalog/enrich workflows.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'Workflows bulk-upsert without custom jobs.',
        'owner': '@elastic/workflows-eng',
        'refs': 'github-catalog-*.yaml',
    },
    {
        'id': 'WF-003',
        'component': 'Workflows',
        'status': 'Partial',
        'priority': 'P2',
        'title': 'Ergonomic checkpoint read (data.loadCheckpoint)',
        'requirement': 'Single step to read sync-state document by id/query; returns cursor fields for while loops.',
        'current': 'Step exists; SDLC package uses verbose elasticsearch.search + data.set chains.',
        'gap': 'High YAML boilerplate; error-prone checkpoint patterns across 25 workflows.',
        'deliverable': 'Document canonical checkpoint recipe; optional step enhancements (default id scheme).',
        'acceptance': 'Refactor one SDLC catalog workflow to loadCheckpoint with fewer steps; package test passes.',
        'owner': '@elastic/workflows-eng',
        'refs': 'docs/data_pipeline/checkpoint_pattern.md',
    },
    {
        'id': 'WF-004',
        'component': 'Workflows',
        'status': 'Partial',
        'priority': 'P2',
        'title': 'elasticsearch.bulk batch ingest ergonomics',
        'requirement': 'Write arrays of docs in one step with stable ids (connector page → bulk).',
        'current': 'Step exists; package uses per-item elasticsearch.index in foreach.',
        'gap': 'Large pages → many steps/executions; slower and harder to rate-limit.',
        'deliverable': 'Bulk ingest cookbook + foreach-in-step or native array body support documented.',
        'acceptance': 'Example workflow indexes 100 project items in one bulk step.',
        'owner': '@elastic/workflows-eng',
        'refs': 'workflow_editor.spec.ts elasticsearch.bulk',
    },
    {
        'id': 'WF-005',
        'component': 'Workflows',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Rate-limit aware retry/backoff step primitive',
        'requirement': 'When connector returns shouldBackoff/rateLimit, workflow pauses with jitter without manual wait step sizing.',
        'current': 'GitHub actions return shouldBackoff; workflows use fixed `wait` duration.',
        'gap': 'Fragile under GitHub rate limits; duplicate wait logic in every GraphQL workflow.',
        'deliverable': 'Optional `retry`/`backoff` on connector steps or dedicated `wait.forRateLimit` step.',
        'acceptance': 'GraphQL catalog workflow removes hardcoded wait; passes rate-limit integration test.',
        'owner': '@elastic/workflows-eng + @elastic/stack-connectors',
        'refs': 'github.ts shouldBackoff; github-catalog-repos.yaml wait step',
    },
    {
        'id': 'WF-006',
        'component': 'Workflows',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Inter-workflow dependency orchestration',
        'requirement': 'Express “run B after A succeeds” or package-level DAG without external cron ordering docs.',
        'current': 'README documents recommended order; no platform enforcement.',
        'gap': 'Mis-ordered enablement causes empty indices; ops burden.',
        'deliverable': 'Workflow `depends_on` or Fleet manifest workflow groups with ordered enablement.',
        'acceptance': 'Enabling epic enrichment auto-checks catalog workflow freshness or chains trigger.',
        'owner': '@elastic/workflows-eng',
        'refs': 'README recommended order',
    },
    {
        'id': 'WF-007',
        'component': 'Workflows',
        'status': 'Partial',
        'priority': 'P2',
        'title': 'Workflow execution history API for ops dashboards',
        'requirement': 'Stable API for last run status, duration, failure reason — consumable from ES|QL view or kibana.request.',
        'current': 'Ingest health dashboard needs workflow run metadata; partial APIs exist.',
        'gap': 'No documented contract for package-level pipeline health monitoring.',
        'deliverable': 'Documented internal API + example ES|QL/workflow for ingest health.',
        'acceptance': 'sdlc_ingest_health ES|QL view joins checkpoint + last workflow run.',
        'owner': '@elastic/workflows-eng',
        'refs': 'sdlc_ingest_health.yml',
    },
    {
        'id': 'WF-008',
        'component': 'Workflows',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Runtime Liquid for connector-id and agent-id',
        'requirement': 'Reference integration policy vars or secrets at execution time (not only install-time substitution).',
        'current': 'Fleet install-time string replace; consts are static after install.',
        'gap': 'Multi-policy or rotated connectors need reinstall or manual edits.',
        'deliverable': 'Liquid/const resolution from integration context OR documented Fleet-only pattern.',
        'acceptance': 'Design doc chooses install-time vs runtime; one path fully supported.',
        'owner': '@elastic/workflows-eng',
        'refs': 'PLATFORM_RFC non-goals vs FLEET-004',
    },
    {
        'id': 'WF-009',
        'component': 'Workflows',
        'status': 'Done',
        'priority': '—',
        'title': 'ai.agent step with structured JSON output',
        'requirement': 'Agent step returns schema-validated object for downstream elasticsearch.index.',
        'current': 'agent-coverage-analysis and agent-scope-alignment workflows persist to sdlc-agent-insights.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'Agent output indexed without custom plugin code.',
        'owner': '@elastic/workflows-eng + @elastic/agent-builder',
        'refs': 'agent-coverage-analysis.yaml',
    },
    {
        'id': 'WF-010',
        'component': 'Workflows',
        'status': 'Partial',
        'priority': 'P3',
        'title': 'ES|QL materialize step for scheduled snapshots',
        'requirement': 'Persist ES|QL view results to snapshot index on schedule from workflow.',
        'current': 'Pattern used in examples; not all deployments expose step consistently.',
        'gap': 'Snapshot/trend indices require author knowledge of advanced step.',
        'deliverable': 'First-class `elasticsearch.esql.materialize` documented for Fleet packages.',
        'acceptance': 'Package example snapshots ES|QL view daily without custom job plugin.',
        'owner': '@elastic/workflows-eng',
        'refs': 'ARCHITECTURE_VISION_GAPS epic snapshot workflow',
    },
    # Connectors v2
    {
        'id': 'CONN-001',
        'component': 'Connectors v2',
        'status': 'Done',
        'priority': '—',
        'title': 'GitHub GraphQL ingest plane (runQueryTemplate, graphqlQuery)',
        'requirement': 'Workflow-only read-only GraphQL with template library for org-scale ingest.',
        'current': 'orgCatalog.*, activity.*, graph.* templates in kbn-connector-specs.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'SDLC GitHub workflows run without custom github step types.',
        'owner': '@elastic/stack-connectors',
        'refs': 'kbn-connector-specs/src/specs/github/',
    },
    {
        'id': 'CONN-002',
        'component': 'Connectors v2',
        'status': 'Done',
        'priority': '—',
        'title': 'Slack ingest actions (isTool: false)',
        'requirement': 'Paginated listUsers, getChannelHistory, getConversationReplies for ETL.',
        'current': 'Implemented in slack connector spec.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'Slack catalog workflows use slack2.* ingest actions.',
        'owner': '@elastic/stack-connectors',
        'refs': 'slack.ts',
    },
    {
        'id': 'CONN-003',
        'component': 'Connectors v2',
        'status': 'Done',
        'priority': '—',
        'title': 'Salesforce SOQL ingest action',
        'requirement': 'Paginated soqlIngest for Case/object ETL with hasMore cursor.',
        'current': 'Implemented; used by salesforce-catalog-cases workflow.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'Salesforce workflow paginates without custom step.',
        'owner': '@elastic/stack-connectors',
        'refs': 'salesforce.ts soqlIngest',
    },
    {
        'id': 'CONN-004',
        'component': 'Connectors v2',
        'status': 'Done',
        'priority': '—',
        'title': 'Google Drive metadata ingest actions',
        'requirement': 'listFilesIngest, parseDriveUrlsFromText, getFileMetadata, parseCommaSeparatedIds.',
        'current': 'Implemented for Phase D design-doc metadata workflows.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'gdrive workflows run with metadata-only ingest.',
        'owner': '@elastic/stack-connectors',
        'refs': 'google_drive.ts',
    },
    {
        'id': 'CONN-005',
        'component': 'Connectors v2',
        'status': 'Partial',
        'priority': 'P2',
        'title': 'isTool:false vs isTool:true contract documentation',
        'requirement': 'Published contributor guide: ingest actions are workflow-only; agent tools are MCP-exposed.',
        'current': 'Enforced in code/tests; not in central connector contributor docs.',
        'gap': 'Integration authors may expose heavy ingest to Agent Builder or omit ingest actions.',
        'deliverable': 'CONTRIBUTOR.md section + lint rule in connector-specs CI.',
        'acceptance': 'New connector PR checklist includes isTool classification.',
        'owner': '@elastic/stack-connectors',
        'refs': 'PLATFORM_RFC Related assets connector contract',
    },
    {
        'id': 'CONN-006',
        'component': 'Connectors v2',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Package-shippable connector query templates',
        'requirement': 'Integrations extend GitHub GraphQL template library from Fleet package without kibana core merge.',
        'current': 'Templates live only in kbn-connector-specs platform repo.',
        'gap': 'Product-specific GraphQL requires platform PR for every new template.',
        'deliverable': 'Manifest-registered read-only templates merged at runtime OR documented extension API.',
        'acceptance': 'Package adds template name callable via runQueryTemplate without core code change.',
        'owner': '@elastic/stack-connectors',
        'refs': 'github/graphql/templates',
    },
    {
        'id': 'CONN-007',
        'component': 'Connectors v2',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Connector preflight health API for integrations',
        'requirement': 'Single API: given connector id + action, verify auth/scopes before enabling 25 workflows.',
        'current': 'Per-connector test in Stack Management only; no batch preflight.',
        'gap': 'Fleet install success does not prove ingest will run.',
        'deliverable': 'Workflow step or Fleet install hook calling connector.test with package-defined matrix.',
        'acceptance': 'SDLC integration policy save runs preflight and surfaces actionable errors.',
        'owner': '@elastic/stack-connectors + @elastic/fleet',
        'refs': 'github test handler',
    },
    {
        'id': 'CONN-008',
        'component': 'Connectors v2',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Normalized bulk-ready connector page action (optional)',
        'requirement': 'Connector action returns ES bulk lines for common ingest patterns (optional alternative to workflow foreach index).',
        'current': 'Raw GraphQL/REST returned; normalization in workflow YAML or omitted.',
        'gap': 'Thinner workflows need shared normalize in platform or stay verbose.',
        'deliverable': 'Optional getProjectV2PageNormalized-style actions delegating to shared libraries.',
        'acceptance': 'Design doc: when to use normalized action vs stock elasticsearch.index loop.',
        'owner': '@elastic/stack-connectors',
        'refs': 'PLATFORM_GAPS_RFC C-1',
    },
    # Agent Builder
    {
        'id': 'AB-001',
        'component': 'Agent Builder',
        'status': 'Done',
        'priority': '—',
        'title': 'AgentBuilderManagementSetup create/update/delete API',
        'requirement': 'Fleet and plugins create persisted agents without HTTP-only routes.',
        'current': 'agent_builder_management_api.ts exposed on setup.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'Fleet stepInstallAgentAssets creates agents programmatically.',
        'owner': '@elastic/agent-builder',
        'refs': 'agent_builder_management_api.ts',
    },
    {
        'id': 'AB-002',
        'component': 'Agent Builder',
        'status': 'Done',
        'priority': '—',
        'title': 'Platform ES|QL tools for package agents',
        'requirement': 'Agents use platform.core.executeEsql / generateEsql without product builtin tools.',
        'current': 'SDLC agents tool allowlist is platform-only.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'No security_solution agent registration required.',
        'owner': '@elastic/agent-builder',
        'refs': 'kibana/agent/sdlc-*.yaml',
    },
    {
        'id': 'AB-003',
        'component': 'Agent Builder',
        'status': 'Partial',
        'priority': 'P2',
        'title': 'Integration knowledge → agent context wiring',
        'requirement': 'Installed package knowledge base automatically available to fleet-installed agents in same integration.',
        'current': 'Knowledge indexed on install; agent context linkage requires Enterprise + manual settings.',
        'gap': 'Agents may not retrieve schema docs unless admin enables integration knowledge globally.',
        'deliverable': 'Document required settings; optional auto-bind knowledge to fleet agents.',
        'acceptance': 'Post-install checklist guarantees agent answers cite sdlc-github-data-model.md.',
        'owner': '@elastic/agent-builder + @elastic/fleet',
        'refs': 'step_save_knowledge_base.ts',
    },
    {
        'id': 'AB-004',
        'component': 'Agent Builder',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Package-managed agents read-only in UI',
        'requirement': 'Fleet-installed agents show managed state; prevent accidental instruction edits.',
        'current': 'Agents editable like user-created agents.',
        'gap': 'Upgrade cannot reliably restore package agent definitions.',
        'deliverable': 'managed_by_package metadata + UI restrictions (mirror workflow managed pattern).',
        'acceptance': 'Edit attempt on fleet agent shows warning or blocks save.',
        'owner': '@elastic/agent-builder',
        'refs': 'PLATFORM_RFC open questions §1',
    },
    {
        'id': 'AB-005',
        'component': 'Agent Builder',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Persisted skills from Fleet package assets',
        'requirement': 'Ship agent skills/tool presets as Fleet assets (not only markdown knowledge base).',
        'current': 'Skills registered in code; package ships YAML agents + markdown only.',
        'gap': 'Complex tool presets require plugin registration.',
        'deliverable': 'KibanaAssetType.skill or agent YAML skill_ids resolution to packaged skill defs.',
        'acceptance': 'Package ships skill without TypeScript plugin registration.',
        'owner': '@elastic/agent-builder + @elastic/fleet',
        'refs': 'PLATFORM_RFC open questions §4',
    },
    {
        'id': 'AB-006',
        'component': 'Agent Builder',
        'status': 'Done',
        'priority': '—',
        'title': 'REPLACE_WITH_FLEET_AGENT_* install substitution',
        'requirement': 'Workflow ai.agent steps reference fleet agent ids resolved at install.',
        'current': 'substituteFleetAgentIds in step_install_workflow_assets.',
        'gap': 'None.',
        'deliverable': '—',
        'acceptance': 'agent workflows run with substituted agent-id after install.',
        'owner': '@elastic/fleet',
        'refs': 'substituteFleetAgentIds',
    },
    # Platform content / docs
    {
        'id': 'DOC-001',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P1',
        'title': 'Fleet package authoring guide — Kibana-only ETL integration',
        'requirement': 'Official doc: manifest vars, workflow/agent assets, ES|QL views, knowledge base, install order, uninstall.',
        'current': 'SDLC README is product-specific; data_pipeline docs are workflow-centric not Fleet-centric.',
        'gap': 'No generic “build a Fleet ETL integration” guide on elastic.co/docs.',
        'deliverable': 'New doc under Fleet + cross-links to workflows/connectors.',
        'acceptance': 'Third team ships minimal integration using guide without SDLC source copy.',
        'owner': '@elastic/fleet + @elastic/docs',
        'refs': 'sdlc_intel_fleet_package/README.md',
    },
    {
        'id': 'DOC-002',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P1',
        'title': 'Workflow ETL cookbook for integration packages',
        'requirement': 'Recipes: checkpoint, pagination, rate limits, bulk index, cross-index enrich, ai.agent persist.',
        'current': 'data_pipeline/README partial; SDLC encodes patterns only in YAML.',
        'gap': 'High YAML complexity without copy-paste templates.',
        'deliverable': 'kbn-workflows/docs/integration_etl/ with Fleet-oriented examples.',
        'acceptance': 'Cookbook covers 5 patterns used by sdlc_intel workflows.',
        'owner': '@elastic/workflows-eng',
        'refs': 'docs/data_pipeline/',
    },
    {
        'id': 'DOC-003',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Placeholder substitution convention specification',
        'requirement': 'Document REPLACE_WITH_* naming, manifest var mapping, agent id convention, multi-value vars.',
        'current': 'Convention implicit in SDLC Fleet installer code.',
        'gap': 'Each integration invents placeholders until FLEET-001 lands.',
        'deliverable': 'Fleet package spec section + JSON schema for vars → placeholders.',
        'acceptance': 'Spec referenced by FLEET-001 implementation and package validator.',
        'owner': '@elastic/fleet',
        'refs': 'step_install_workflow_assets.ts',
    },
    {
        'id': 'DOC-004',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P2',
        'title': 'Agent Builder fleet agent asset authoring guide',
        'requirement': 'How to author kibana/agent/*.yaml: tools, instructions, connector placeholders, knowledge deps.',
        'current': 'PLATFORM_RFC is design doc only; no user-facing authoring guide.',
        'gap': 'Integration teams cannot ship agents without reading Fleet source.',
        'deliverable': 'Agent Builder docs section “Package-managed agents”.',
        'acceptance': 'Guide lists required YAML fields matching parseFleetAgentYaml validation.',
        'owner': '@elastic/agent-builder + @elastic/docs',
        'refs': 'step_install_agent_assets.ts parseFleetAgentYaml',
    },
    {
        'id': 'DOC-005',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P2',
        'title': 'GitHub action connector vs content connector decision guide',
        'requirement': 'When to use .github workflow ingest vs Elasticsearch GitHub content connector; dual-pipeline pattern.',
        'current': 'No platform doc; confusion risk for integration authors.',
        'gap': 'Teams may choose wrong connector model for Projects V2 vs repo corpus.',
        'deliverable': 'Connectors docs decision matrix (workflow ETL vs search sync).',
        'acceptance': 'Doc explicitly states Projects V2 requires action connector GraphQL plane.',
        'owner': '@elastic/stack-connectors + @elastic/docs',
        'refs': 'es-connectors-github vs kbn-connector-specs github',
    },
    {
        'id': 'DOC-006',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Integration alerting templates enablement guide',
        'requirement': 'How alerting_rule_template assets work today vs after FLEET-002; enablement UX.',
        'current': 'SDLC README describes desired UX; platform behavior partial.',
        'gap': 'Admins unclear why templates do not auto-create rules.',
        'deliverable': 'Alerting + Fleet doc updated when FLEET-002 ships; interim “manual enable” doc now.',
        'acceptance': 'Doc matches actual install behavior per release.',
        'owner': '@elastic/alerting + @elastic/fleet',
        'refs': 'kibana/alerting_rule_template/',
    },
    {
        'id': 'DOC-007',
        'component': 'Platform content',
        'status': 'Open',
        'priority': 'P3',
        'title': 'Reference architecture: packaged multi-source ETL on Kibana',
        'requirement': 'Single architecture diagram: Fleet assets + connectors + workflows + agents + ES indices.',
        'current': 'Split across SDLC README, data_pipeline RFC, PLATFORM_RFC.',
        'gap': 'No one platform architecture page for integration authors.',
        'deliverable': 'Platform architecture doc (product-agnostic) on elastic.co/docs.',
        'acceptance': 'Diagram labels only generic components; SDLC cited as example appendix.',
        'owner': '@elastic/fleet + @elastic/workflows-eng',
        'refs': 'Multiple RFCs',
    },
]


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_gap_table(doc: Document, gaps: list[dict], title: str) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['ID', 'Component', 'Priority', 'Status', 'Title', 'Owner']
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], 'D9E2F3')
    for gap in gaps:
        row = table.add_row().cells
        row[0].text = gap['id']
        row[1].text = gap['component']
        row[2].text = gap['priority']
        row[3].text = gap['status']
        row[4].text = gap['title']
        row[5].text = gap['owner']
    doc.add_paragraph()


def add_ticket_section(doc: Document, gap: dict) -> None:
    doc.add_heading(f"{gap['id']}: {gap['title']}", level=3)
    fields = [
        ('Component', gap['component']),
        ('Status', gap['status']),
        ('Priority', gap['priority']),
        ('Owner', gap['owner']),
        ('References', gap['refs']),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)

    for label, key in [
        ('Requirement', 'requirement'),
        ('Current state', 'current'),
        ('Gap', 'gap'),
        ('Proposed platform deliverable', 'deliverable'),
        ('Acceptance criteria', 'acceptance'),
    ]:
        p = doc.add_paragraph()
        p.add_run(f'{label}\n').bold = True
        p.add_run(gap[key])

    doc.add_paragraph('Suggested ticket title:')
    doc.add_paragraph(f"[{gap['component']}] {gap['title']}", style='Intense Quote')

    doc.add_paragraph('Suggested PR scope:')
    bullets = gap['deliverable'].split(';')
    for b in bullets:
        if b.strip() and b.strip() != '—':
            doc.add_paragraph(b.strip(), style='List Bullet')
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    title = doc.add_heading('Platform gaps & requirements', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Packaged Kibana ETL integrations (SDLC Intelligence reference)')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'Version 1.0 · {date.today().isoformat()}\n')
    meta.add_run('Audience: @elastic/fleet, @elastic/workflows-eng, @elastic/stack-connectors, @elastic/agent-builder\n')
    meta.add_run('Scope: Generic platform capabilities only (not SDLC product logic)')

    doc.add_page_break()

    # Executive summary
    doc.add_heading('1. Executive summary', level=1)
    doc.add_paragraph(
        'The SDLC Intelligence Fleet package is the reference implementation of a Kibana-only integration: '
        'Elasticsearch index templates, ES|QL views, dashboards, 25+ ingest workflows, persisted Agent Builder agents, '
        'integration knowledge base, and alerting rule templates — with no Elastic Agent.'
    )

    open_gaps = [g for g in GAPS if g['status'] == 'Open']
    partial = [g for g in GAPS if g['status'] == 'Partial']
    done = [g for g in GAPS if g['status'] == 'Done']

    by_component: dict[str, dict[str, int]] = {}
    for g in GAPS:
        comp = g['component']
        by_component.setdefault(comp, {'Open': 0, 'Partial': 0, 'Done': 0})
        by_component[comp][g['status']] += 1

    summary = doc.add_table(rows=1, cols=5)
    summary.style = 'Table Grid'
    sh = summary.rows[0].cells
    for i, h in enumerate(['Metric', 'Open', 'Partial', 'Done', 'Total']):
        sh[i].text = h
        set_cell_shading(sh[i], 'D9E2F3')

    rows = [
        ('All gaps', len(open_gaps), len(partial), len(done), len(GAPS)),
    ]
    for comp, counts in sorted(by_component.items()):
        total = sum(counts.values())
        rows.append((comp, counts['Open'], counts['Partial'], counts['Done'], total))

    for row_data in rows:
        cells = summary.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)

    doc.add_paragraph()
    doc.add_paragraph(
        'Each gap ID maps 1:1 to a GitHub issue or PR. Sections 4–8 provide ticket-ready detail blocks. '
        'Platform content gaps (DOC-*) are counted explicitly — documentation is a required platform deliverable.'
    )

    # Scope
    doc.add_heading('2. Scope & methodology', level=1)
    doc.add_paragraph('In scope:', style='List Bullet')
    for item in [
        'Fleet EPM Kibana/Elasticsearch asset install pipeline',
        'Workflows stock-step ETL primitives used by packaged workflows',
        'Connectors v2 (kbn-connector-specs) ingest actions with isTool: false',
        'Agent Builder persisted agents installed via Fleet',
        'Platform documentation and authoring guides for integration teams',
    ]:
        doc.add_paragraph(item, style='List Bullet 2')

    doc.add_paragraph('Out of scope:', style='List Bullet')
    for item in [
        'SDLC domain logic (epic gates, team taxonomy, Salesforce field semantics)',
        'SDLC-specific TypeScript job handlers or security_solution plugin code',
        'GitHub content connector as replacement for Projects V2 ingest (see DOC-005)',
        'Elastic Agent-based ingest (this integration is Kibana-only)',
    ]:
        doc.add_paragraph(item, style='List Bullet 2')

    # Reference model
    doc.add_heading('3. Reference platform model', level=1)
    doc.add_paragraph(
        'A packaged Kibana ETL integration (exemplified by sdlc_intel) requires the following generic platform stack:'
    )
    model = doc.add_table(rows=1, cols=3)
    model.style = 'Table Grid'
    mh = model.rows[0].cells
    for i, h in enumerate(['Layer', 'Platform owner', 'Generic capability']):
        mh[i].text = h
        set_cell_shading(mh[i], 'D9E2F3')

    model_rows = [
        ('Assets & bootstrap', 'Fleet', 'Index templates, ILM, ES|QL views, dashboards, knowledge base, workflow/agent YAML install'),
        ('Policy binding', 'Fleet', 'Manifest vars → placeholder substitution; optional enable-on-install'),
        ('Fetch', 'Connectors v2', 'isTool:false ingest actions; paginated API/GraphQL; rate-limit metadata'),
        ('Orchestrate', 'Workflows', 'Scheduled stock-step ETL: search/index/bulk, foreach/while, ai.agent'),
        ('Analyze', 'Agent Builder', 'Package-managed persisted agents; platform ES|QL tools; integration knowledge'),
        ('Operate', 'Fleet + Workflows', 'Uninstall cleanup, upgrade reconcile, ingest health, alerting templates'),
        ('Author', 'Platform content', 'Guides, cookbooks, conventions for integration teams'),
    ]
    for layer, owner, cap in model_rows:
        r = model.add_row().cells
        r[0].text = layer
        r[1].text = owner
        r[2].text = cap

    doc.add_paragraph()

    # Master register
    doc.add_heading('4. Master gap register', level=1)
    add_gap_table(doc, GAPS, '4.1 All gaps')

    open_p1 = [g for g in GAPS if g['status'] == 'Open' and g['priority'] == 'P1']
    add_gap_table(doc, open_p1, '4.2 Open P1 gaps (implement first)')

    # Detailed sections by component
    doc.add_page_break()
    doc.add_heading('5. Ticket-ready gap specifications', level=1)
    doc.add_paragraph(
        'Copy each subsection into a GitHub issue. Suggested labels: platform, fleet-package, integration-etl.'
    )

    components = ['Fleet', 'Workflows', 'Connectors v2', 'Agent Builder', 'Platform content']
    for comp in components:
        doc.add_heading(comp, level=2)
        for gap in GAPS:
            if gap['component'] == comp:
                add_ticket_section(doc, gap)

    # PR batching
    doc.add_heading('6. Suggested PR batching', level=1)
    batches = [
        ('PR-A: Fleet substitution generalization', ['FLEET-001', 'DOC-003']),
        ('PR-B: Integration alerting materialization', ['FLEET-002', 'DOC-006']),
        ('PR-C: Post-install ops UX', ['FLEET-004', 'FLEET-005', 'FLEET-013']),
        ('PR-D: Workflow ETL ergonomics', ['WF-003', 'WF-004', 'WF-005', 'DOC-002']),
        ('PR-E: Rate limits & orchestration', ['WF-005', 'WF-006', 'CONN-005']),
        ('PR-F: Managed assets UI', ['FLEET-006', 'AB-004', 'FLEET-012']),
        ('PR-G: Platform docs pack', ['DOC-001', 'DOC-004', 'DOC-005', 'DOC-007']),
    ]
    batch_table = doc.add_table(rows=1, cols=2)
    batch_table.style = 'Table Grid'
    bh = batch_table.rows[0].cells
    bh[0].text = 'PR theme'
    bh[1].text = 'Gap IDs'
    set_cell_shading(bh[0], 'D9E2F3')
    set_cell_shading(bh[1], 'D9E2F3')
    for theme, ids in batches:
        row = batch_table.add_row().cells
        row[0].text = theme
        row[1].text = ', '.join(ids)

    doc.add_heading('7. References', level=1)
    refs = [
        'x-pack/solutions/security/packages/sdlc_intel_fleet_package/README.md',
        'x-pack/solutions/security/packages/sdlc_intel_fleet_package/docs/PLATFORM_RFC_KIBANA_ASSET_EXTENSIONS.md',
        'src/platform/packages/shared/kbn-workflows/docs/data_pipeline/README.md',
        'x-pack/platform/plugins/shared/fleet/server/services/epm/packages/install_state_machine/steps/step_install_workflow_assets.ts',
        'x-pack/platform/plugins/shared/fleet/server/services/epm/packages/install_state_machine/steps/step_install_agent_assets.ts',
        'x-pack/platform/plugins/shared/fleet/server/services/epm/packages/install_state_machine/steps/step_create_alerting_assets.ts',
        'src/platform/packages/shared/kbn-connector-specs/src/specs/github/github.ts',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f'Wrote {OUTPUT}')


if __name__ == '__main__':
    main()
